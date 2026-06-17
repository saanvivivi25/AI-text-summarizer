import google.generativeai as genai
import sqlite3
from datetime import datetime
from docx import Document

# GEMINI API CONFIGURATION

genai.configure(api_key="YOUR_API_KEY_HERE")
model = genai.GenerativeModel("gemini-2.5-flash")


# DATABASE FUNCTIONS

def create_database():
    try:
        conn = sqlite3.connect("summaries.db")
        cursor = conn.cursor()

        cursor.execute("""
        CREATE TABLE IF NOT EXISTS Summaries (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            original_text TEXT,
            summary TEXT,
            key_points TEXT,
            action_items TEXT,
            created_at TEXT
        )
        """)

        conn.commit()
        conn.close()

    except Exception as e:
        print("Database Error:", e)


def save_to_database(original_text, ai_response):
    try:
        conn = sqlite3.connect("summaries.db")
        cursor = conn.cursor()

        created_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        cursor.execute("""
        INSERT INTO Summaries
        (original_text, summary, key_points, action_items, created_at)
        VALUES (?, ?, ?, ?, ?)
        """, (
            original_text,
            ai_response,
            "",
            "",
            created_at
        ))

        conn.commit()
        conn.close()

    except Exception as e:
        print("Database Error:", e)


def view_database():
    try:
        conn = sqlite3.connect("summaries.db")
        cursor = conn.cursor()

        cursor.execute("SELECT * FROM Summaries")
        records = cursor.fetchall()

        if not records:
            print("\nDatabase is empty.")
        else:
            print("\n===== SUMMARIES DATABASE =====")

            for row in records:
                print(f"\nID: {row[0]}")
                print(f"Original Text:\n{row[1]}")
                print(f"\nAI Response:\n{row[2]}")
                print(f"\nKey Points: {row[3]}")
                print(f"Action Items: {row[4]}")
                print(f"Created At: {row[5]}")
                print("-" * 60)

        conn.close()

    except Exception as e:
        print("Error:", e)


def delete_summary():
    try:
        summary_id = int(input("Enter Summary ID to delete: "))

        conn = sqlite3.connect("summaries.db")
        cursor = conn.cursor()

        cursor.execute(
            "SELECT * FROM Summaries WHERE id = ?",
            (summary_id,)
        )

        record = cursor.fetchone()

        if record:
            cursor.execute(
                "DELETE FROM Summaries WHERE id = ?",
                (summary_id,)
            )
            conn.commit()
            print("Summary deleted successfully.")
        else:
            print("Summary ID not found.")

        conn.close()

    except ValueError:
        print("Please enter a valid ID.")
    except Exception as e:
        print("Error:", e)

# SEARCH SUMMARY BY ID 
def search_summary_by_id():
    try:
        summary_id = int(input("Enter Summary ID: "))

        conn = sqlite3.connect("summaries.db")
        cursor = conn.cursor()

        cursor.execute(
            "SELECT * FROM Summaries WHERE id = ?",
            (summary_id,)
        )

        record = cursor.fetchone()

        if record:
            print("\n===== SUMMARY FOUND =====")
            print(f"ID: {record[0]}")
            print(f"Original Text:\n{record[1]}")
            print(f"\nAI Response:\n{record[2]}")
            print(f"\nKey Points: {record[3]}")
            print(f"Action Items: {record[4]}")
            print(f"Created At: {record[5]}")
        else:
            print("No summary found with that ID.")

        conn.close()

    except ValueError:
        print("Please enter a valid numeric ID.")
    except Exception as e:
        print("Error:", e)

# EXPORT TO WORD

def export_to_word():
    try:
        conn = sqlite3.connect("summaries.db")
        cursor = conn.cursor()

        cursor.execute("SELECT * FROM Summaries")
        records = cursor.fetchall()

        if not records:
            print("No summaries available to export.")
            conn.close()
            return

        document = Document()
        document.add_heading(
            "AI Text Summaries Report",
            level=1
        )

        for row in records:
            document.add_heading(
                f"Summary ID: {row[0]}",
                level=2
            )

            document.add_paragraph(
                f"Original Text:\n{row[1]}"
            )

            document.add_paragraph(
                f"AI Response:\n{row[2]}"
            )

            document.add_paragraph(
                f"Key Points:\n{row[3]}"
            )

            document.add_paragraph(
                f"Action Items:\n{row[4]}"
            )

            document.add_paragraph(
                f"Created At: {row[5]}"
            )

            document.add_paragraph(
                "-" * 50
            )

        document.save("summary_report.docx")

        conn.close()

        print("\nReport exported successfully.")
        print("File Name: summary_report.docx")

    except Exception as e:
        print("Error:", e)

# GENERATE SUMMARY

def generate_summary(text):
    try:
        prompt = f"""
Analyze the following text and provide:

1. Summary
2. Key Points
3. Action Items

Text:
{text}
"""

        response = model.generate_content(prompt)
        return response.text

    except Exception as e:
        return f"Error: {e}"

# MAIN PROGRAM

def main():
    create_database()

    while True:
        print("\n===== AI TEXT SUMMARIZER =====")
        print("1. Generate Summary")
        print("2. View Database")
        print("3. Export to Word")
        print("4. Search Summary by ID")
        print("5. Delete Summary")
        print("6. Exit")

        choice = input("Enter Choice: ")

        if choice == "1":
            article = input("\nEnter article:\n")

            result = generate_summary(article)

            print("\n===== GENERATED REPORT =====")
            print(result)

            save_to_database(
                article,
                result
            )

            print("\nSummary saved to database.")

        elif choice == "2":
            view_database()

        elif choice == "3":
            export_to_word()

        elif choice == "4":
            search_summary_by_id()

        elif choice == "5":
            delete_summary()

        elif choice == "6":
            print("Thank you!")
            break

        else:
            print("Invalid choice. Try again.")


# RUN PROGRAM

if __name__ == "__main__":
    main()